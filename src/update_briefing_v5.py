#!/usr/bin/env python3
"""Update the ST briefing .docx to v5.

Adds the agent-injection measurements, the two flow diagrams, and the Work IQ /
Fabric IQ gap analyses. Renumbers the references section accordingly.

Word holds an exclusive lock on the file while it is open, so this builds to a
temporary path and only then attempts the copy, falling back to a sibling name
with a clear message rather than losing the work.
"""
import copy
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

DOCX = Path("/mnt/c/Users/jpontvianne/Documents/Azure/ST/IQs/"
            "Foundry-Work-Web-IQ-Security-Briefing.docx")
TMP = Path("/tmp/briefing-v5.docx")
IMG_DIR = Path("/tmp")

ACCENT = RGBColor(0x0A, 0x68, 0xC1)
WARN = RGBColor(0xC1, 0x28, 0x0A)


# --------------------------------------------------------------- helpers ----
class Writer:
    """Appends content immediately before a given anchor paragraph."""

    def __init__(self, doc, anchor):
        self.doc = doc
        self.anchor = anchor

    def _insert(self, item):
        # lxml elements with no children are falsy, so test for None explicitly.
        el = getattr(item, "_p", None)
        if el is None:
            el = getattr(item, "_tbl", None)
        self.anchor._p.addprevious(el)
        return item

    def h(self, text, level=1):
        return self._insert(self.doc.add_heading(text, level=level))

    def p(self, text="", bold=False, italic=False, size=None, color=None,
          style=None, space_after=6):
        para = self.doc.add_paragraph(style=style)
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        para.paragraph_format.space_after = Pt(space_after)
        return self._insert(para)

    def bullet(self, text, bold_lead=None):
        para = self.doc.add_paragraph(style="List Bullet")
        if bold_lead:
            para.add_run(bold_lead).bold = True
        para.add_run(text)
        para.paragraph_format.space_after = Pt(3)
        return self._insert(para)

    def numbered(self, text, bold_lead=None):
        para = self.doc.add_paragraph(style="List Number")
        if bold_lead:
            para.add_run(bold_lead).bold = True
        para.add_run(text)
        para.paragraph_format.space_after = Pt(3)
        return self._insert(para)

    def image(self, path, width_in=6.6, caption=None):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(path), width=Inches(width_in))
        self._insert(para)
        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.italic = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            cap.paragraph_format.space_after = Pt(10)
            self._insert(cap)

    def table(self, headers, rows, widths=None):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, htext in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(htext)
            run.bold = True
            run.font.size = Pt(9)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                para = cells[i].paragraphs[0]
                # **bold** segments
                for j, seg in enumerate(str(val).split("**")):
                    if seg:
                        run = para.add_run(seg)
                        run.bold = (j % 2 == 1)
                        run.font.size = Pt(8.5)
        if widths:
            for r in t.rows:
                for i, w in enumerate(widths):
                    r.cells[i].width = Inches(w)
        self._insert(t)
        self._insert(self.doc.add_paragraph())
        return t


def find_heading(doc, prefix):
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip().startswith(prefix):
            return para
    return None


# ------------------------------------------------------------------ main ----
def main():
    if not DOCX.exists():
        sys.exit(f"missing {DOCX}")
    for img in ("flow-foundry-iq.png", "flow-web-iq.png"):
        if not (IMG_DIR / img).exists():
            sys.exit(f"missing diagram {IMG_DIR / img} - run src/render_flow_diagrams.py")

    doc = Document(str(DOCX))

    refs = find_heading(doc, "8.")
    if refs is None:
        sys.exit("could not find the references heading to insert before")
    refs.text = "10.  Key references (official)"

    w = Writer(doc, refs)

    # ================================================================= §7b ===
    w.h("7b.  Agent VNet injection — what we injected and what it proved", 1)

    w.p("The first round of POC testing configured a delegated subnet but never "
        "placed an agent in it, so the claim that agent egress can be brought "
        "inside ST's network was only half-tested. That gap has now been closed: "
        "an agent was genuinely injected, and the flows below were measured "
        "rather than drawn.", italic=True)

    w.h("7b.1  Delegation is not injection", 2)
    w.p("Subnet delegation only reserves a subnet. The proof that the platform "
        "has actually claimed it is a serviceAssociationLink. After injection, "
        "snet-agent carries:")
    w.p("delegations: Microsoft.App/environments      "
        "serviceAssociationLinks: legionservicelink      "
        "ipConfigurations: 0      networkInterfaces: none",
        italic=True, size=9, color=ACCENT)
    w.p("Two practical consequences for ST:")
    w.bullet("There is no NIC and no IP address to reference in a firewall rule "
             "or a CMDB entry. The subnet itself is the only handle.",
             bold_lead="The agent runtime is serverless. ")
    w.bullet("was applied by PATCHing an existing Foundry account. It is not a "
             "create-time-only property, contrary to the usual assumption.",
             bold_lead="Injection ")

    w.h("7b.2  Flow A — agent to Foundry IQ", 2)
    w.p("Measured run: status completed, tool azure_ai_search, 1074 prompt / 157 "
        "completion tokens. The answer quoted a canary string planted only in the "
        "private index, so retrieval genuinely occurred rather than the model "
        "answering from memory.")
    w.image(IMG_DIR / "flow-foundry-iq.png", 6.7,
            "Figure 2 — Agent to Foundry IQ. Every hop is private; no public IP "
            "address appears anywhere in the path.")

    w.h("7b.3  Flow B — agent to Web IQ", 2)
    w.p("Measured run: status completed, tool bing_grounding, 4204 prompt / 97 "
        "completion tokens, returning a correctly cited answer about ST's "
        "quarterly revenue. The agent was created and executed from inside the VNet.")
    w.image(IMG_DIR / "flow-web-iq.png", 6.7,
            "Figure 3 — Agent to Web IQ. The Bing call is made service-side by "
            "Microsoft and never enters ST's network.")

    # ================================================================= §7c ===
    w.h("7c.  The decisive test — can an NSG separate Foundry IQ from Web IQ?", 1)

    w.p("This is the measurement most relevant to ST's governance and proxy "
        "questions, and its result is not what the architecture diagrams would "
        "lead you to expect.")
    w.p("Three configurations were applied to the injected subnet and both agents "
        "re-run each time:")

    w.table(
        ["#", "Configuration on snet-agent", "Foundry IQ agent", "Web IQ agent"],
        [["1", "Baseline — unrestricted egress", "completed", "completed"],
         ["2", "Deny * → Internet (Allow AzureCloud)", "completed", "**completed**"],
         ["3", "Deny * → 10.30.2.10/32 (Search PE)", "**failed**", "completed"]],
        widths=[0.4, 3.1, 1.5, 1.5])

    w.p("Run 3 is the control, and it is what makes run 2 interpretable. Because a "
        "serverless injected subnet has no NIC, it is entirely reasonable to "
        "suspect that NSG rules there are accepted by ARM and then never "
        "enforced. They are enforced: denying the Azure AI Search private "
        "endpoint broke the Foundry IQ agent with tool_user_error: "
        "search_access_error.")
    w.p("Given that the NSG demonstrably works, run 2 is decisive. Blocking all "
        "Internet egress from the agent subnet did not stop Grounding with Bing. "
        "The Bing call therefore does not originate from ST's network at all.",
        bold=True)

    w.h("7c.1  What this means for ST", 2)
    w.bullet("Retrieval traffic really does traverse ST's subnet, so NSGs, UDRs, "
             "forced tunnelling and a corporate proxy all apply to it normally.",
             bold_lead="Foundry IQ is network-governable. ")
    w.bullet("There is no packet for ST to intercept. It cannot be proxied, "
             "inspected, or blocked by any firewall rule ST is able to write. "
             "Blocking api.bing.microsoft.com at ST's own egress accomplishes "
             "nothing, because the call is not made from ST's egress.",
             bold_lead="Web IQ is not. ")
    w.bullet("yes for Foundry IQ, no for Web IQ.",
             bold_lead="This is the honest answer to \"can we add a proxy?\" — ")

    w.p("The only effective controls for Web IQ are therefore administrative:")
    w.numbered("Do not create a Grounding with Bing connection on the project.")
    w.numbered("Deny Microsoft.Bing/accounts creation with Azure Policy.")
    w.numbered("Restrict who may add project connections (Azure AI Developer / "
               "Owner RBAC on the Foundry project).")
    w.numbered("Alert on connection-creation events in the Foundry account "
               "activity log.")

    w.h("7c.2  Correction to the earlier network analysis", 2)
    w.p("An earlier draft attributed agent egress to the customer's NAT Gateway. "
        "That was wrong and is corrected here: snet-agent has no NAT Gateway and "
        "no route table attached at all. The NAT address 4.165.97.252 belongs to "
        "snet-app, where the test client VM runs.", italic=True)

    w.h("7c.3  Side-by-side summary", 2)
    w.table(
        ["Question", "Foundry IQ", "Web IQ"],
        [["Does traffic traverse the customer VNet?", "Yes (proved by control test)", "**No**"],
         ["Can an NSG / UDR / firewall control it?", "Yes", "No"],
         ["Can a proxy intercept it?", "Yes", "No"],
         ["Public IP anywhere in the path?", "No", "Yes, service-side"],
         ["Inside the Azure DPA boundary?", "Yes", "**No** — First Party Consumption Service"],
         ["Data residency", "Sweden Central", "location: global"],
         ["Primary control surface", "Network + RBAC", "Connection governance + Azure Policy"]],
        widths=[2.6, 2.0, 2.0])

    # ================================================================== §8 ===
    w.h("8.  Work IQ — what is missing, and how to fulfil it", 1)

    w.p("Work IQ could not be exercised in the POC tenant. Rather than record "
        "that as \"not tested\", each prerequisite was probed until the service "
        "itself stated the reason.", italic=True)

    w.table(
        ["Probe", "Result"],
        [["GET /v1.0/subscribedSkus", "one SKU only: AAD_PREMIUM_P2"],
         ["GET /v1.0/sites/root", "**BadRequest: Tenant does not have a SPO license.**"],
         ["GET /v1.0/external/connections", "Unauthenticated (no Graph permission consented)"]],
        widths=[2.8, 3.8])

    w.p("The POC tenant is a bare Azure tenant: it has Entra ID but no Microsoft "
        "365 workloads. Work IQ retrieves from the Microsoft 365 substrate — "
        "SharePoint, OneDrive, Exchange, Teams and Graph connector content — so "
        "with no substrate there is nothing to retrieve and no Azure-side "
        "configuration can compensate.")
    w.p("This is a limitation of the POC tenant, not of the product. ST's "
        "production tenant already has the substrate, so most of the list below "
        "is \"confirm\" rather than \"acquire\".", italic=True)

    w.h("8.1  Requirements in dependency order", 2)
    w.table(
        ["#", "Requirement", "Provided by", "ST's likely status"],
        [["1", "M365 tenant with SharePoint Online", "M365 licensing", "**Already have**"],
         ["2", "Content in the substrate (SPO / OneDrive / Exchange / Teams)", "ST business units", "Already have"],
         ["3", "Microsoft 365 Copilot licences per querying user", "M365 add-on SKU", "**Verify — usually the gap**"],
         ["4", "Graph connectors for non-M365 content", "ST IT", "Optional"],
         ["5", "Semantic index over the tenant", "automatic after (3)", "Follows from (3)"],
         ["6", "Entra app registration with Graph scopes", "ST identity team", "To do"],
         ["7", "Admin consent for those scopes", "Global Administrator", "To do"],
         ["8", "Work IQ / M365 Copilot Retrieval API or MCP endpoint", "preview enrolment", "**Verify availability**"]],
        widths=[0.35, 2.7, 1.6, 1.85])

    w.p("Item 3 is what most often stalls a pilot. Work IQ honours per-user "
        "licensing, so an unlicensed caller receives no results — and receives "
        "them silently, as an empty answer rather than an error. Budget time for "
        "that confusing failure mode.")

    w.h("8.2  Remediation steps and who performs them", 2)
    w.bullet("confirm an M365 E3/E5 base SKU plus the Microsoft_365_Copilot "
             "add-on in the M365 admin centre, or via Graph subscribedSkus.",
             bold_lead="M365 administrator — ")
    w.bullet("assign Copilot licences to the pilot group only.",
             bold_lead="M365 administrator — ")
    w.bullet("register the calling application and request the delegated Graph "
             "scopes the use case genuinely needs (Files.Read.All, "
             "Sites.Read.All, Mail.Read, ExternalItem.Read.All).",
             bold_lead="ST identity team — ")
    w.bullet("grant admin consent. This is tenant-wide and will need ST's change "
             "process, so start it early; it is usually the long pole.",
             bold_lead="Global Administrator — ")
    w.bullet("optionally configure Graph connectors, verifying that each source's "
             "ACLs are mapped to Entra identities, otherwise permission trimming "
             "silently fails.", bold_lead="ST IT — ")
    w.bullet("call the endpoint with a delegated user token and confirm that two "
             "users with different SharePoint permissions receive different "
             "results. That single test exercises the entire security model.",
             bold_lead="Validation — ")

    w.h("8.3  Privacy points to raise in the meeting", 2)
    w.bullet("Work IQ never returns content the calling user cannot already open "
             "in SharePoint. This is the strongest privacy statement available "
             "for any of the four IQ services.",
             bold_lead="Permission trimming is per-user. ")
    w.bullet("In the Foundry IQ tests, retrieval used a single managed identity, "
             "so every end user collapsed into one object ID in the logs and all "
             "users saw the same index content. Work IQ is the opposite. If ST "
             "needs per-user authorisation and per-user audit, Work IQ provides "
             "it natively while Foundry IQ requires it to be built in the client "
             "layer.", bold_lead="Contrast with the Foundry IQ finding. ")
    w.bullet("Work IQ activity is recorded in the Microsoft 365 unified audit log, "
             "not the Log Analytics workspace used by Foundry IQ and AI Search. "
             "ST will operate two separate audit surfaces, typically owned by two "
             "different teams.", bold_lead="Audit lands in Purview. ")
    w.bullet("Work IQ operates under the M365 Data Protection Addendum and the EU "
             "Data Boundary, not the Azure DPA. Both are acceptable, but they are "
             "different contracts.", bold_lead="Different data boundary. ")
    w.bullet("If ST uses application permissions instead of delegated ones, "
             "per-user trimming is bypassed entirely. Insist on the delegated "
             "flow for anything touching human-readable content.",
             bold_lead="App-only tokens defeat the model. ")
    w.bullet("Work IQ indexes in place and creates no new copy, so ST's existing "
             "SharePoint retention, sensitivity labels and DLP continue to apply "
             "unchanged.", bold_lead="Retention is unchanged. ")

    # ================================================================== §9 ===
    w.h("9.  Fabric IQ — what is missing, and how to fulfil it", 1)

    w.p("The ARM error alone is misleading and sent the first investigation down "
        "the wrong path:", italic=True)
    w.p("PUT .../Microsoft.Fabric/capacities/... → 401 \"Unable to authorize with "
        "Azure Active Directory.\"", italic=True, size=9, color=WARN)
    w.p("The Fabric data plane gives the real answer:")
    w.p("GET https://api.fabric.microsoft.com/v1/capacities → 401 "
        "{\"errorCode\":\"UserNotLicensed\"}", italic=True, size=9, color=WARN)

    w.table(
        ["Probe", "Result"],
        [["Microsoft.Fabric RP registration", "Registered — so the RP is not the problem"],
         ["Existing capacities in the subscription", "none"],
         ["ARM PUT F2 capacity", "401 Unable to authorize with Azure Active Directory"],
         ["Fabric data plane /v1/capacities", "**401 UserNotLicensed**"],
         ["Power BI /v1.0/myorg/groups", "404 Not Found"]],
        widths=[2.8, 3.8])

    w.p("Root cause: no Fabric / Power BI tenant has ever been initialised in "
        "this Entra tenant. Microsoft.Fabric/capacities is an ARM resource, but "
        "the resource provider delegates authorisation to the Power BI service; "
        "with no Power BI tenant object, that check fails and surfaces as a "
        "generic ARM 401. Registering the resource provider is not enough, and "
        "neither is subscription Owner.", bold=True)

    w.h("9.1  Requirements in dependency order", 2)
    w.table(
        ["#", "Requirement", "Provided by", "Notes"],
        [["1", "Fabric / Power BI tenant initialised in Entra", "first licensed interactive sign-in", "**The blocker here**"],
         ["2", "A Fabric (free or Pro) licence on one user", "M365 licensing", "Creates (1)"],
         ["3", "Microsoft.Fabric RP registered", "Azure subscription", "already done"],
         ["4", "An F-SKU capacity (F2 suffices)", "Azure subscription", "blocked by (1)"],
         ["5", "A workspace bound to that capacity", "Fabric admin", ""],
         ["6", "Tenant settings: Fabric items, Ontology/Graph preview, data agents", "**Fabric Admin portal only**", "not scriptable"],
         ["7", "A lakehouse with data, plus an ontology/graph", "ST data team", "the actual work"],
         ["8", "A Fabric data agent published over the ontology", "ST data team", ""],
         ["9", "A project connection from Foundry to the data agent", "Azure AI Developer", ""]],
        widths=[0.35, 2.8, 1.75, 1.6])

    w.h("9.2  Remediation steps", 2)
    w.bullet("assign any Fabric-capable licence (the free Fabric licence is "
             "sufficient) to one user and have them sign in once interactively at "
             "app.fabric.microsoft.com. The first successful sign-in creates the "
             "tenant object. UserNotLicensed must disappear before anything else "
             "will work.", bold_lead="Unblock — ")
    w.bullet("create the F2 capacity. It costs roughly €0.36/hour and can be "
             "paused when idle, so pause it between workshops rather than "
             "deleting it.", bold_lead="Then — ")
    w.bullet("in the Fabric Admin portal, enable \"Users can create Fabric "
             "items\", the Ontology / Graph preview switches, and \"Users can "
             "create and use Fabric data agents\" — each scoped to a security "
             "group rather than the whole organisation. There is no supported API "
             "for these switches, so this step cannot be automated and needs a "
             "named owner and a change window.", bold_lead="Fabric administrator — ")
    w.bullet("create the workspace and lakehouse, load data, then define the "
             "ontology / graph. This is the substance of Fabric IQ: it is a "
             "modelled semantic layer, not a connector, and it is where the "
             "effort actually sits.", bold_lead="ST data team — ")
    w.bullet("have two users with different OneLake permissions ask the same "
             "question through the data agent and confirm the answers differ.",
             bold_lead="Validation — ")

    w.h("9.3  Networking and privacy points", 2)
    w.bullet("Fabric is a SaaS service reached over its own public endpoints, and "
             "a Foundry agent calling a Fabric data agent is a service-to-service "
             "call — closer to the Web IQ pattern than the private endpoint "
             "pattern. Expect Fabric IQ to be governed by Fabric tenant settings "
             "and workspace RBAC rather than by network controls. Private Link "
             "for Fabric exists but is a separate tenant-level feature that must "
             "be assessed on its own; do not assume the Foundry IQ private "
             "posture extends to it.",
             bold_lead="Fabric IQ is not the Foundry IQ flow. ")
    w.bullet("Fabric data agent activity lands in the M365 unified audit log "
             "(Purview), not Log Analytics. Combined with Work IQ, ST will have "
             "both Azure Monitor and Purview to govern.",
             bold_lead="Two audit surfaces again. ")
    w.bullet("A Fabric data agent enforces the calling user's OneLake permissions "
             "when invoked with user identity — again the opposite of the Foundry "
             "IQ managed-identity behaviour measured in the POC.",
             bold_lead="Per-user authorisation. ")
    w.bullet("The agent's language model follows Fabric Copilot settings, "
             "including the \"data leaves your geography\" toggle. Confirm this "
             "explicitly with ST; the default is not always what an EU-resident "
             "customer expects.", bold_lead="Cross-geography processing. ")
    w.bullet("Ontology, graph and data agents are in preview. Preview terms "
             "exclude the standard SLA and may differ on data handling; this must "
             "be stated plainly before ST puts real engineering data into a pilot.",
             bold_lead="Preview status. ")

    doc.save(str(TMP))
    print(f"  built {TMP}")

    try:
        shutil.copyfile(TMP, DOCX)
        print(f"  updated {DOCX}")
    except PermissionError:
        alt = DOCX.with_name(DOCX.stem + "-v5.docx")
        shutil.copyfile(TMP, alt)
        print(f"  {DOCX.name} is locked by Word; wrote {alt.name} instead")


if __name__ == "__main__":
    main()
