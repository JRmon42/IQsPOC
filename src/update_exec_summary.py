#!/usr/bin/env python3
"""Surface the POC's decisive result in the executive summary.

The measurements landed in sections 7b/7c, but the executive summary still
described a posture in which VNet injection implies network control over
everything an agent does. That is the specific belief the POC disproved, so a
reader who stops after section 0 would take away the wrong answer to ST's
proxy question. This inserts a short, plainly worded correction there.
"""
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

DOCX = Path("/mnt/c/Users/jpontvianne/Documents/Azure/ST/IQs/"
            "Foundry-Work-Web-IQ-Security-Briefing.docx")
TMP = Path("/tmp/briefing-v5b.docx")

ACCENT = RGBColor(0x0A, 0x68, 0xC1)


def main():
    doc = Document(str(DOCX))

    anchor = None
    started = False
    for para in doc.paragraphs:
        if para.style.name == "Heading 1" and para.text.strip().startswith("0."):
            started = True
            continue
        if started and para.text.strip().startswith("The single most important message"):
            anchor = para
            break
    if anchor is None:
        raise SystemExit("could not locate the executive summary paragraph")

    if "proxy" in (anchor._p.getnext().text or "" if anchor._p.getnext() is not None else ""):
        print("  addendum already present; nothing to do")
        return

    def add(text, bold=False, italic=False, lead=None, size=None):
        para = doc.add_paragraph()
        if lead:
            run = para.add_run(lead)
            run.bold = True
            if size:
                run.font.size = Pt(size)
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        para.paragraph_format.space_after = Pt(5)
        anchor._p.addnext(para._p)
        return para

    # addnext inserts directly after the anchor, so build the block in reverse.
    add("Full detail and reproduction steps are in sections 7b and 7c.",
        italic=True, size=9)
    add("only by not creating the Bing connection, and enforcing that with "
        "Azure Policy.",
        lead="Web IQ can therefore be governed ")
    add("no for Web IQ. Denying all internet egress from the agent subnet did "
        "not stop Grounding with Bing, while denying the Azure AI Search "
        "private endpoint immediately broke Foundry IQ retrieval. The second "
        "result proves the network rules are genuinely enforced, which is what "
        "makes the first one conclusive: the Bing call is made service-side by "
        "Microsoft and never enters the customer network, so there is nothing "
        "to proxy, inspect or firewall.",
        lead="On the proxy question specifically: yes for Foundry IQ, ")
    add("These statements were subsequently tested in a working deployment "
        "rather than taken from documentation. One result changed the "
        "recommended posture and is worth reading before the meeting.",
        italic=True)

    doc.save(str(TMP))
    shutil.copyfile(TMP, DOCX)
    print(f"  executive summary updated in {DOCX.name}")


if __name__ == "__main__":
    main()
